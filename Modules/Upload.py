##### Imports 
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
from pathlib import Path
import os
import docxtpl
import jinja2
import docx 
import time


##### Functions
def parse_filesystem(filesystem):
    filesystemIndex = []
    os.chdir(filesystem)
    
    allYears = os.listdir()
    for year in allYears:
        currentYear = str(year)
        os.chdir(currentYear)
        allMonths = os.listdir()
        for month in allMonths:
            currentMonth = str(month)
            os.chdir(currentMonth)
            allDays = os.listdir()
            for day in allDays:
                currentDay = str(day)
                os.chdir(currentDay)
                allEvents = os.listdir()
                for event in allEvents:
                    currentEvent = str(event)
                    
                    images = os.listdir(currentEvent)
                    location = filesystem + currentYear + "/" + currentMonth + "/" + currentDay + "/" + currentEvent + "/" 
                    filesystemIndex.append([currentYear, currentMonth, currentDay, currentEvent, images, location])

                os.chdir(str(Path().resolve().parent))
            os.chdir(str(Path().resolve().parent))
        os.chdir(str(Path().resolve().parent))
    os.chdir(str(Path().resolve().parent))
  
    #Returns a nested list pointing to individual events
    return filesystemIndex

def create_folder(parent_folder_id, subfolder_name):
    newFolder = drive.CreateFile({'title': subfolder_name, "parents": [{"kind": "drive#fileLink", "id": \
    parent_folder_id}],"mimeType": "application/vnd.google-apps.folder"})
    newFolder.Upload()
    return newFolder

def verify_root(rootFileName):
    filesystemIDs = {}
    fileList = drive.ListFile({'q': "'root' in parents and trashed=false"}).GetList()
    flag = False
    test = []
    
    for file in fileList:
        if file['mimeType'] == "application/vnd.google-apps.folder":
            filesystemIDs[file['title']] = file['id']
            flag = True
    
    if(not flag):
        print("Sorry! no file found....")
    
    if rootFileName in filesystemIDs:
        print("root file exists with ID: ", filesystemIDs[rootFileName])
    else:
        tmp = drive.CreateFile({"title":rootFileName,"mimeType":folder})
        tmp.Upload()
        filesystemIDs[rootFileName] = tmp["id"]
        print("root file created with ID: ", filesystemIDs[rootFileName])
    
    return filesystemIDs

def generate_filesystem(filesystemIndex):

    for files in filesystemIndex:
    
        year = files[0]
        yearMonth = year + "/" + files[1]
        yearMonthDay = yearMonth + "/" + files[2]
        eventName = files[3]
        yearMonthDayEvent = yearMonthDay + "/" + eventName
        images = files[4]
        location = files[5]

        if year not in filesystemIDs:
            filesystemIDs[year] = create_folder(filesystemIDs[filesystem], year)["id"]
            print("Created Folder: ", year)

        if yearMonth not in filesystemIDs:
            filesystemIDs[yearMonth] = create_folder(filesystemIDs[year], yearMonth)["id"]
            print("Created Folder: ", yearMonth)
            
        if yearMonthDay not in filesystemIDs:
            filesystemIDs[yearMonthDay] = create_folder(filesystemIDs[yearMonth], yearMonthDay)["id"]
            print("Created Folder: ", yearMonthDay)

        eventNameFolder = create_folder(filesystemIDs[yearMonthDay], eventName)
        filesystemIDs[yearMonthDayEvent] = eventNameFolder["id"]
        
        eventImageFolder = create_folder(filesystemIDs[yearMonthDayEvent], "Images")
        filesystemIDs[yearMonthDayEvent + "/images"] = eventImageFolder["id"]

        
        permission = eventNameFolder.InsertPermission({'type': 'anyone','value': 'anyone','role': 'reader'})
        
        #SHARABLE LINK
        link=eventNameFolder['alternateLink']
        files.append(link)
        
        for i in images:
            os.chdir(location)
            tmpfile = drive.CreateFile({"name":i, "parents": [{"kind": "drive#fileLink", "id": eventImageFolder["id"]}]})
            tmpfile.SetContentFile(i)
            tmpfile.Upload()
            print("Uploaded ", i, " to ", link)

def clean_data(data):
    cleandata = []
    for x in data:
        date = x[0]+"/"+x[1]+"/"+x[2]+"/"
        event = x[3]
    
        fileID = filesystemIDs[date+event]
        fileLink = x[6]
    
        cleandata.append([date,event,fileID,fileLink])
    return cleandata         
            
def upload_file(file, fileID):
    tmpfile = drive.CreateFile({"name":fileID, "parents": [{"kind": "drive#fileLink", "id": fileID}]})
    tmpfile.SetContentFile(file)
    tmpfile.Upload()
    print("Uploaded ", fileID, " to ", fileID)
    tmpfile.InsertPermission({'type': 'anyone','value': 'anyone','role': 'reader'})
    link = tmpfile["alternateLink"]
    return link
    
def generate_summary_files(cleandata):
    os.chdir(filesystem)
    for x in cleandata:

        doc = docxtpl.DocxTemplate(r"C:\Users\tyler\Desktop\Workspace\TTT Admin Suite\Template.docx")
        filename = str(x[2])+".docx" 
    
        link = docxtpl.RichText()
        link.add("Filesystem", url_id=doc.build_url_id(x[3]), color='#0000FF')
    
        context = {'DATE' : x[0]}
        context["FILESYSTEM"] = link 
        context["EVENT"] = x[1]
    
        doc.render(context)
        doc.save(filename) 

        link = upload_file(filesystem + filename, x[2])
        x.append(link)
    return cleandata
   
def generate_index_templates(cleandata):
    
    months = []
    years = []

    for x in cleandata:
        event = x[1]
        summary = x[4]
        archive = x[3]  
        
        date = x[0]
        date = date.split("/") 

        month = date[1]
        year = date[0]
        
        dateindex = ("DATE" + str(year + month) + event).replace(" ", "")
        eventindex = ("EVENT" + str(year + month) + event).replace(" ", "")
        
        filename = year+month+".docx"
        
        x.append(dateindex) #x[5]
        x.append(eventindex) #[6]
        x.append(filename) #x[7]
        
        
        if years == []:
            years = [year]
    
        if year not in years:
            years = [year]
            months = []

        if month in months:
            doc = docx.Document(index+filename)
            para = doc.add_paragraph("{{r "+ dateindex +" }} {{r "+ eventindex +" }}")
            doc.save(index+filename)
        
        else:
            months.append(month)

            doc = docx.Document(rootfilesystem+"Template2.docx")
            para = doc.add_paragraph("{{r "+ dateindex +" }} {{r "+ eventindex +" }}")
            doc.save(index+filename)


def generate_index_files(cleandata):
    generate_index_templates(cleandata)
    for x in cleandata:     
        event = x[1]
        summary = x[4]
        archive = x[3]   
        
        date = x[0]
        date.split("/")
        
        month = date[1]
        year = date[0]
        
        dateindex = x[5]
        eventindex = x[6]

        filename = index+x[7]
        
        doc = docxtpl.DocxTemplate(filename)
    
        linkrt = docxtpl.RichText()
        linkrt.add(event, url_id = doc.build_url_id(summary), color="#0000FF")
    
        datert = docxtpl.RichText()
        datert.add(date, color='#FFFFFF')
        
        context[dateindex] = datert
        context[eventindex] = linkrt
    
        doc.save(filename)
        time.sleep(.25)
    
    
    for x in os.listdir(index):
        doc = docxtpl.DocxTemplate(index+x)
    
        doc.render(context)
        doc.save(index+x)


##### Variables
drive = GoogleDrive(GoogleAuth())
filesystem = "C:/Users/tyler/Documents/GitHub/The-Trialists-Toolkit/Archives/Filesystem/"
rootfilesystem = "C:/Users/tyler/Desktop/Workspace/TTT Admin Suite/"
index = filesystem + r"index/"
os.makedirs(index)
folder = "application/vnd.google-apps.folder"
filesystemIndex = parse_filesystem(filesystem)
filesystemIDs = verify_root(filesystem)
context = {}

##### Main
generate_filesystem(filesystemIndex)
print("\n[ FILESYSTEM SUCCESSFULLY GENERATED ]\n")

cleandata = generate_summary_files(clean_data(filesystemIndex))
print("\n[ SUMMARY FILES SUCCESSFULLY GENERATED ]\n")

generate_index_files(cleandata)
print("\n[ INDEX FILES SUCCESSFULLY GENERATED ]\n")

for x in os.listdir(index):
    fileID = filesystemIDs[ x[0:4] + r"/" + x[4:6] ]
    
    upload_file(index+x,fileID)



print("\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n",
    "#####PROGRAM EXECUTION FIN. DEBUG OUTPUT BEGINNING.#####"
)

#fix mnth and yr variables